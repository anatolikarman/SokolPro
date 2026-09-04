"""DOCX export mirroring NotesWordExporter.java, block for block."""
from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor

from models import Client, Session
from notes_html import Block, LineBreak, TextRun, parse_blocks


def build(client: Client, sessions: List[Session]) -> bytes:
    doc = Document()

    _add_title(doc, f"Заметки клиента: {client.name}", 16)
    _add_subtitle(doc, "Все сеансы с заметками, от старых к новым")

    has_backstory = bool(client.backstory and client.backstory.strip())
    if has_backstory:
        _append_html(doc, client.backstory)
        _add_title(doc, "Конец предыстории", 14)

    if not has_backstory and not sessions:
        doc.add_paragraph("Заметок нет")

    for session in sessions:
        _add_title(doc, session.formatted_date, 14)
        _append_html(doc, session.notes)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_title(doc: Document, text: str, font_size: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)


def _add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor.from_string("777777")
    run.font.size = Pt(10)


def _append_html(doc: Document, html: str) -> None:
    for block in parse_blocks(html):
        _append_block(doc, block)


def _append_block(doc: Document, block: Block) -> None:
    if block.kind == "title1":
        _add_title(doc, block.text, 15)
    elif block.kind == "title2":
        _add_title(doc, block.text, 13)
    elif block.kind == "title3":
        _add_title(doc, block.text, 12)
    elif block.kind == "bullet":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.add_run("• ")
        _write_runs(p, block.runs)
    else:
        p = doc.add_paragraph()
        _write_runs(p, block.runs)


def _write_runs(paragraph, runs) -> None:
    for run in runs:
        if isinstance(run, LineBreak):
            paragraph.add_run().add_break()
        elif isinstance(run, TextRun):
            r = paragraph.add_run(run.text)
            r.bold = run.bold
            r.italic = run.italic
            if run.underline:
                r.underline = True
